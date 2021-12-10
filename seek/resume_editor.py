from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import Pt


analyst_resume_path = '/Users/tony/Documents/programming/job-search/resume/analyst/'

def main():
    document = Document(analyst_resume_path + 'resume.docx')
    tables = document.tables
    document.tables[0].style.font.name = 'arial'
    document.tables[0].style.font.size = Pt(24)
    table = tables[0]
    name = table.cell(0,0)
    contact = table.cell(0,1)
    main = table.cell(1,0)
    skills_and_certs = table.cell(1,1)
    p1 = locate('Python', skills_and_certs)
    skills_and_certs.add_paragraph
    document.save(analyst_resume_path + 'resume0.docx')

def locate(text, doc):
    for p in doc.paragraphs:
        if p.text == text:
            return p

def copy_style(t, r):
    r.style = t.style
    r.font.name = t.font.name
    r.font.color.rgb = t.font.color.rgb
    r.font.color.theme_color = t.font.color.theme_color
    r.font.size = t.font.size
    

def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

if __name__ == '__main__':
    main()
